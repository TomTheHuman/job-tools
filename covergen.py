'''Template cover letter generator

Lightweight tool for quickly replacing company and position information in a
cover letter template to speed up job application submission time.
'''

from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import pyperclip

# Replacement keys - this can be changed but make sure it matches the key
# in your template document and is unique enough to not conflict with other text
COMPANY_KEY = '[Target Company]'
POSITION_KEY = '[Target Position]'

# ex. 'Template.docx'
TEMPLATE_FILE = 'Cover Letter - Template.docx' 
# ex. 'C:/Users/me/Documents/'
TEMPLATE_DIR = 'C:/Users/grimm/OneDrive/Documents/Professional/' 
# ex. 'C:/Users/me/Documents/Letters/'
DESTINATION_DIR = 'C:/Users/grimm/OneDrive/Documents/Professional/Cover Letters/' 

# Font configurations - docx will overwrite the font and text size. Adjust this
# to match the styling set in your template
FONT_NAME = 'Segoe UI'
FONT_SIZE = Pt(11)

if __name__ == '__main__':
    # Collect user input
    company = input('What is the company name?\n')
    position = input('What is the position?\n')
    clipboard = input('Copy text to clipboard?\n')
    need_pdf = input('Would you like a PDF copy? (y/n)\n')
    file_name = 'Cover Letter - ' + company + '.docx'

    # Get template document and set styles
    document = Document(TEMPLATE_DIR + TEMPLATE_FILE)
    style = document.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE

    text_cell = None

    # Find cell with text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                  if COMPANY_KEY in paragraph.text:
                      text_cell = cell

    count = 0
    for paragraph in text_cell.paragraphs:
        if COMPANY_KEY in paragraph.text:
            count += 1
            paragraph.text = paragraph.text.replace(COMPANY_KEY, company)
            paragraph.style = document.styles['Normal']
    print(f'{COMPANY_KEY} replaced with {company}: {count}')

    # Replace position references
    count = 0
    for paragraph in text_cell.paragraphs:
        if POSITION_KEY in paragraph.text:
            count += 1
            paragraph.text = paragraph.text.replace(POSITION_KEY, position)
            paragraph.style = document.styles['Normal']
    print(f'{POSITION_KEY} replaced with {position}: {count}')

    # Save document
    new_file = DESTINATION_DIR + file_name
    document.save(new_file)

    # Optional - save PDF copy
    if need_pdf == 'y':
        convert(new_file)
        print('PDF copy generated')

    # Optional - copy text to clipboard
    if clipboard == 'y':
        pyperclip.copy(text_cell.text)
        print('Cover letter text copied to clipboard')

    print('All done!')